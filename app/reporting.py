"""Reporting and export utilities for VoiceTracer."""

import json
from dataclasses import asdict
from io import BytesIO
import datetime
from typing import Dict, List, Any, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

from app.analysis import AnalysisResult, MetricResult

class ReportGenerator:
    """Generates various report formats from AnalysisResult."""

    @staticmethod
    def to_dict(analysis: AnalysisResult) -> Dict[str, Any]:
        """Convert AnalysisResult (and nested dataclasses) to a dictionary."""
        return asdict(analysis)

    @classmethod
    def generate_json(cls, analysis: AnalysisResult) -> bytes:
        """Generate a JSON report."""
        data = cls.to_dict(analysis)
        data["generated_at"] = datetime.datetime.now().isoformat()
        return json.dumps(data, indent=4).encode("utf-8")

    @classmethod
    def generate_docx(
        cls, 
        analysis: AnalysisResult, 
        sections_to_include: List[str], 
        statement: str,
        negotiated_text: Optional[str] = None,
        images: Optional[Dict[str, bytes]] = None,
        include_documentation: bool = False
    ) -> bytes:
        """Generate a Word document report."""
        if not Document:
            return b"Error: python-docx not installed."
            
        doc = Document()
        
        # Header
        doc.add_heading("VoiceTracer Analysis Report", 0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

        # Executive Summary
        if "Executive Summary" in sections_to_include:
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(f"Overall Voice Preservation Score: {analysis.score:.2f}/100")
            doc.add_paragraph(f"Classification: {analysis.classification}")
            doc.add_paragraph(f"Consistency Score: {analysis.consistency_score:.2f}")

        # Data Only Section
        doc.add_heading("Detailed Analysis Data", level=1)
        
        # Component Breakdown Table
        doc.add_heading("Component Breakdown", level=2)
        comp_table = doc.add_table(rows=1, cols=2)
        comp_table.style = 'Table Grid'
        comp_table.rows[0].cells[0].text = "Component"
        comp_table.rows[0].cells[1].text = "Score"
        
        components = [
            ("Authenticity Markers", analysis.components.get("Authenticity", 0.0)),
            ("Lexical Identity", analysis.components.get("Lexical", 0.0)),
            ("Structural Identity", analysis.components.get("Structural", 0.0)),
            ("Stylistic Identity", analysis.components.get("Stylistic", 0.0)),
            ("Voice Consistency", analysis.consistency_score),
        ]
        for name, val in components:
            row = comp_table.add_row().cells
            row[0].text = name
            row[1].text = f"{val:.2f}"

        # Quick Stats Table
        doc.add_heading("Quick Stats", level=2)
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Table Grid'
        stats_table.rows[0].cells[0].text = "Metric"
        stats_table.rows[0].cells[1].text = "Value"
        
        stats = [
            ("Word Delta", str(analysis.word_delta)),
            ("Sentence Delta", str(analysis.sentence_delta)),
            ("AI-isms Total", str(analysis.ai_ism_total)),
            ("Original Word Count", str(analysis.original_word_count)),
            ("Rewrite Word Count", str(analysis.edited_word_count)),
            ("Original Sentence Count", str(analysis.original_sentence_count)),
            ("Rewrite Sentence Count", str(analysis.edited_sentence_count)),
        ]
        for name, val in stats:
            row = stats_table.add_row().cells
            row[0].text = name
            row[1].text = val

        # Similarity Index
        doc.add_heading("Similarity Index", level=2)
        sim_table = doc.add_table(rows=1, cols=2)
        sim_table.style = 'Table Grid'
        sim_table.rows[0].cells[0].text = "Metric"
        sim_table.rows[0].cells[1].text = "Score"
        
        sims = [
            ("Source vs Rewrite (Cosine)", f"{analysis.source_similarity_cosine * 100:.1f}%"),
            ("Source vs Rewrite (N-gram)", f"{analysis.source_similarity_ngram * 100:.1f}%"),
            ("AI vs Rewrite (Cosine)", f"{analysis.ai_similarity_cosine * 100:.1f}%"),
            ("AI vs Rewrite (N-gram)", f"{analysis.ai_similarity_ngram * 100:.1f}%"),
        ]
        for name, val in sims:
            row = sim_table.add_row().cells
            row[0].text = name
            row[1].text = val

        # Visualizations
        if "Visualizations" in sections_to_include and images:
            doc.add_heading("Visualizations", level=1)
            order = ["Voice Identity Spectrum", "Sentence Length Variation", "Metric Standards", "AI Source AI-isms", "Writer Rewrite AI-isms"]
            for name in order:
                if name in images:
                    doc.add_heading(name, level=2)
                    img_stream = BytesIO(images[name])
                    doc.add_picture(img_stream, width=Inches(5.5))

        # Full Metric Analysis
        if "Full Metric Analysis" in sections_to_include:
            doc.add_heading("Full Metric Deep-Dive", level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'AI Source'
            hdr_cells[2].text = 'Rewrite'
            hdr_cells[3].text = 'Verdict'

            for label, original_val in analysis.metrics_original.items():
                edited_val = analysis.metrics_edited.get(label, 0.0)
                verdict = "N/A"
                for res in analysis.metric_results.values():
                    if res.name == label:
                        verdict = res.verdict
                        break
                
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = f"{original_val:.4f}"
                row_cells[2].text = f"{edited_val:.4f}"
                row_cells[3].text = verdict

        # Repair Preview Decisions
        if "Repair Preview Decisions" in sections_to_include and negotiated_text:
            doc.add_heading("Repair Preview Decisions", level=1)
            doc.add_heading("Final Negotiated Rewrite", level=2)
            doc.add_paragraph(negotiated_text)

        # Authorship Statement
        if "Authorship Documentation Statement" in sections_to_include:
            doc.add_heading("Authorship Statement", level=1)
            doc.add_paragraph(statement)

        # Complete Documentation Appendix
        if include_documentation:
            doc.add_page_break()
            cls._append_docx_documentation(doc, analysis)

        # Save to Stream
        target = BytesIO()
        doc.save(target)
        return target.getvalue()

    @classmethod
    def generate_pdf(
        cls, 
        analysis: AnalysisResult, 
        sections_to_include: List[str], 
        statement: str,
        negotiated_text: Optional[str] = None,
        images: Optional[Dict[str, bytes]] = None,
        include_documentation: bool = False
    ) -> bytes:
        """Generate a PDF document report using fpdf2."""
        if not FPDF:
            return b"Error: fpdf2 not installed."

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Title
        pdf.set_font("Helvetica", "B", 24)
        pdf.cell(0, 20, "VoiceTracer Analysis Report", 0, 1, "C")
        
        # Date
        pdf.set_font("Helvetica", "I", 10)
        pdf.cell(0, 10, f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1, "R")
        pdf.ln(10)

        # Executive Summary
        if "Executive Summary" in sections_to_include:
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, "Executive Summary", 0, 1)
            pdf.set_font("Helvetica", "", 12)
            pdf.cell(0, 10, f"Overall Voice Preservation Score: {analysis.score:.2f}/100", 0, 1)
            pdf.cell(0, 10, f"Classification: {analysis.classification}", 0, 1)
            pdf.cell(0, 10, f"Consistency Score: {analysis.consistency_score:.2f}", 0, 1)
            pdf.ln(5)

        # Data Only Sections
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "Detailed Analysis Data", 0, 1)
        pdf.ln(2)

        # Component Breakdown
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Component Breakdown", 0, 1)
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(100, 10, "Component", 1)
        pdf.cell(40, 10, "Score", 1)
        pdf.ln()
        pdf.set_font("Helvetica", "", 10)
        components = [
            ("Authenticity Markers", analysis.components.get("Authenticity", 0.0)),
            ("Lexical Identity", analysis.components.get("Lexical", 0.0)),
            ("Structural Identity", analysis.components.get("Structural", 0.0)),
            ("Stylistic Identity", analysis.components.get("Stylistic", 0.0)),
            ("Voice Consistency", analysis.consistency_score),
        ]
        for name, val in components:
            pdf.cell(100, 10, name, 1)
            pdf.cell(40, 10, f"{val:.2f}", 1)
            pdf.ln()
        pdf.ln(5)

        # Quick Stats
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Quick Stats", 0, 1)
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(100, 10, "Metric", 1)
        pdf.cell(40, 10, "Value", 1)
        pdf.ln()
        pdf.set_font("Helvetica", "", 10)
        stats = [
            ("Word Delta", str(analysis.word_delta)),
            ("Sentence Delta", str(analysis.sentence_delta)),
            ("AI-isms Total", str(analysis.ai_ism_total)),
            ("Original Word Count", str(analysis.original_word_count)),
            ("Rewrite Word Count", str(analysis.edited_word_count)),
            ("Original Sentence Count", str(analysis.original_sentence_count)),
            ("Rewrite Sentence Count", str(analysis.edited_sentence_count)),
        ]
        for name, val in stats:
            pdf.cell(100, 10, name, 1)
            pdf.cell(40, 10, val, 1)
            pdf.ln()
        pdf.ln(5)

        # Similarity Index
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Similarity Index", 0, 1)
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(100, 10, "Metric", 1)
        pdf.cell(40, 10, "Score", 1)
        pdf.ln()
        pdf.set_font("Helvetica", "", 10)
        sims = [
            ("Source vs Rewrite (Cosine)", f"{analysis.source_similarity_cosine * 100:.1f}%"),
            ("Source vs Rewrite (N-gram)", f"{analysis.source_similarity_ngram * 100:.1f}%"),
            ("AI vs Rewrite (Cosine)", f"{analysis.ai_similarity_cosine * 100:.1f}%"),
            ("AI vs Rewrite (N-gram)", f"{analysis.ai_similarity_ngram * 100:.1f}%"),
        ]
        for name, val in sims:
            pdf.cell(100, 10, name, 1)
            pdf.cell(40, 10, val, 1)
            pdf.ln()
        pdf.ln(5)

        # Visualizations
        if "Visualizations" in sections_to_include and images:
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, "Visualizations", 0, 1)
            order = ["Voice Identity Spectrum", "Sentence Length Variation", "Metric Standards", "AI Source AI-isms", "Writer Rewrite AI-isms"]
            for name in order:
                if name in images:
                    pdf.set_font("Helvetica", "B", 12)
                    pdf.cell(0, 10, name, 0, 1)
                    img_stream = BytesIO(images[name])
                    pdf.image(img_stream, x=10, w=180) 
                    pdf.ln(5)

        # Full Metric Analysis
        if "Full Metric Analysis" in sections_to_include:
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, "Detailed Metric Deep-Dive", 0, 1)
            
            # Table Header
            pdf.set_font("Helvetica", "B", 10)
            col_widths = [60, 30, 30, 40]
            headers = ["Metric", "AI Source", "Rewrite", "Verdict"]
            for i in range(len(headers)):
                pdf.cell(col_widths[i], 10, headers[i], 1)
            pdf.ln()

            # Table Rows
            pdf.set_font("Helvetica", "", 10)
            for label, original_val in analysis.metrics_original.items():
                edited_val = analysis.metrics_edited.get(label, 0.0)
                verdict = "N/A"
                for res in analysis.metric_results.values():
                    if res.name == label:
                        verdict = res.verdict
                        break
                
                pdf.cell(col_widths[0], 10, label, 1)
                pdf.cell(col_widths[1], 10, f"{original_val:.3f}", 1)
                pdf.cell(col_widths[2], 10, f"{edited_val:.3f}", 1)
                pdf.cell(col_widths[3], 10, verdict, 1)
                pdf.ln()
            pdf.ln(5)

        # Repair Preview Decisions
        if "Repair Preview Decisions" in sections_to_include and negotiated_text:
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, "Repair Preview Decisions", 0, 1)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 10, "Final Negotiated Rewrite:", 0, 1)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 10, negotiated_text)
            pdf.ln(5)

        # Authorship Statement
        if "Authorship Documentation Statement" in sections_to_include:
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, "Authorship Statement", 0, 1)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 8, statement)

        # Documentation Appendix
        if include_documentation:
            pdf.add_page()
            cls._append_pdf_documentation(pdf, analysis)

        return bytes(pdf.output())

    @classmethod
    def _get_metric_data(cls, analysis: AnalysisResult) -> List[Dict[str, Any]]:
        """Consolidated metric data for documentation using ASCII math for maximum compatibility."""
        res_map = {r.name: r for r in analysis.metric_results.values()}
        
        m_data = [
            {
                "id": "ai_ism_likelihood",
                "name": "AI-ism Likelihood",
                "purpose": "Detects formulaic AI-like phrases.",
                "calc_plain": "Sums weighted AI-typical phrases (e.g., 'moreover') and divides by total word count.",
                "formula": "L = (Sum(w_i * c_i) / N) * 100",
                "example": f"Found {res_map['AI-ism Likelihood'].details.get('total_weighted_count', 0)} weighted points in your rewrite ({analysis.edited_word_count} words).",
                "human": 3.1,
                "ai": 78.5,
                "meaning": "Higher values indicate more generic, AI-style prose."
            },
            {
                "id": "burstiness",
                "name": "Burstiness",
                "purpose": "Measures sentence length variation and rhythm.",
                "calc_plain": "Standard deviation of sentence lengths divided by the mean sentence length.",
                "formula": "B = StdDev / Mean",
                "example": f"Your sentences varied from {min(res_map['Burstiness'].details.get('sentence_lengths', [0]))} to {max(res_map['Burstiness'].details.get('sentence_lengths', [0]))} words.",
                "human": 1.23,
                "ai": 0.78,
                "meaning": "Higher values show natural human rhythm; lower values are typically robotic."
            },
            {
                "id": "discourse_marker_density",
                "name": "Discourse Marker Density",
                "purpose": "Measures use of structural signposting markers.",
                "calc_plain": "Total transition words (e.g., 'therefore') per 100 words.",
                "formula": "D = (Sum(Markers) / N) * 100",
                "example": f"Detected {res_map['Discourse Marker Density'].details.get('total_markers', 0)} markers in your {analysis.edited_word_count}-word rewrite.",
                "human": 8.0,
                "ai": 18.0,
                "meaning": "AI over-uses markers for clarity; human density is usually moderate."
            },
            {
                "id": "epistemic_hedging",
                "name": "Epistemic Hedging",
                "purpose": "Measures caution and nuance in claims.",
                "calc_plain": "Excess hedging markers (weighted against certainty) per 100 words.",
                "formula": "H = (max(0, hedges - 0.5 * certainty) / N) * 100",
                "example": f"Your rewrite used cautious phrasing to balance specificity.",
                "human": 0.09,
                "ai": 0.04,
                "meaning": "High hedging reflects organic human nuance."
            },
            {
                "id": "function_word_ratio",
                "name": "Function Word Ratio",
                "purpose": "Measures density of grammatical connector words.",
                "calc_plain": "Total count of function words (DT, IN, PRP) over total words.",
                "formula": "F = Sum(Function Words) / N",
                "example": f"Found {res_map['Function Word Ratio'].details.get('function_words', 0)} function words in your text.",
                "human": 0.50,
                "ai": 0.60,
                "meaning": "AI tends to use more 'glue' words to maintain formal structure."
            },
            {
                "id": "information_density",
                "name": "Information Density",
                "purpose": "Concentration of content substance and specificity.",
                "calc_plain": "Weighted sum of content word ratio (70%) and proper noun ratio (30%).",
                "formula": "I = 0.7 * (C/N) + 0.3 * (P/N)",
                "example": f"Your rewrite contains {res_map['Information Density'].details.get('content_words', 0)} content words and {res_map['Information Density'].details.get('proper_nouns', 0)} proper nouns.",
                "human": 0.58,
                "ai": 0.42,
                "meaning": "Higher values show dense, specific human-authored substance."
            },
            {
                "id": "lexical_diversity",
                "name": "Lexical Diversity",
                "purpose": "Vocabulary richness and variety.",
                "calc_plain": "Unique words divided by total words (TTR) or MTLD algorithm.",
                "formula": "TTR = V / N",
                "example": f"Used {res_map['Lexical Diversity'].details.get('unique_words', 0)} unique words out of {analysis.edited_word_count}.",
                "human": 0.55,
                "ai": 0.42,
                "meaning": "Higher diversity indicates a broader and more natural human vocabulary."
            },
            {
                "id": "syntactic_complexity",
                "name": "Syntactic Complexity",
                "purpose": "Measures structural depth and clause use.",
                "calc_plain": "Weighted sum of subordination ratio (60%) and length factor (40%).",
                "formula": "S = 0.6 * R_sub + 0.4 * min(Mean/30, 1)",
                "example": f"Subordination ratio of {res_map['Syntactic Complexity'].details.get('subordination_ratio', 0)} with avg length {analysis.edited_word_count/max(1, analysis.edited_sentence_count):.2f}.",
                "human": 0.54,
                "ai": 0.64,
                "meaning": "AI often over-complicates syntax to sound formal."
            }
        ]
        return sorted(m_data, key=lambda x: x['name'])

    @classmethod
    def _append_docx_documentation(cls, doc, analysis):
        doc.add_heading("Complete Metric Documentation", level=1)
        doc.add_paragraph("Technical breakdown and explanation of the VoiceTracer scoring framework.")
        
        metrics = cls._get_metric_data(analysis)
        for m in metrics:
            doc.add_heading(m['name'], level=2)
            doc.add_paragraph(f"Definition/Purpose: ").add_run(m['purpose']).bold = True
            doc.add_paragraph(f"Plain-Text Calculation: {m['calc_plain']}")
            doc.add_paragraph(f"Example Calculation: {m['example']}")
            doc.add_paragraph(f"Formula (LaTeX): {m['formula']}")
            doc.add_paragraph(f"Current Value Measured: {analysis.metrics_edited.get(m['name'], 0.0):.4f}")
            doc.add_paragraph(f"Ideal Human Standard: {m['human']} | AI Default: {m['ai']}")
            p = doc.add_paragraph(f"Interpretation: ")
            p.add_run(m['meaning']).italic = True

    @classmethod
    def _append_pdf_documentation(cls, pdf, analysis):
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 15, "Complete Metric Documentation", 0, 1)
        pdf.set_font("Helvetica", "I", 10)
        pdf.cell(0, 8, "Technical breakdown of the VoiceTracer scoring framework.", 0, 1)
        pdf.ln(5)

        metrics = cls._get_metric_data(analysis)
        for m in metrics:
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, m['name'], 0, 1)
            
            pdf.set_font("Helvetica", "B", 10)
            pdf.write(5, "Definition: ")
            pdf.set_font("Helvetica", "", 10)
            pdf.write(5, f"{m['purpose']}\n")
            
            pdf.set_font("Helvetica", "B", 10)
            pdf.write(5, "Calculation: ")
            pdf.set_font("Helvetica", "", 10)
            pdf.write(5, f"{m['calc_plain']}\n")
            
            pdf.set_font("Helvetica", "B", 10)
            pdf.write(5, "Example: ")
            pdf.set_font("Helvetica", "", 10)
            pdf.write(5, f"{m['example']}\n")
            
            pdf.set_font("Helvetica", "B", 10)
            pdf.write(5, "Formula: ")
            pdf.set_font("Helvetica", "", 10)
            pdf.write(5, f"{m['formula']}\n")
            
            pdf.set_font("Helvetica", "B", 10)
            pdf.write(5, "Current Value: ")
            pdf.set_font("Helvetica", "", 10)
            pdf.write(5, f"{analysis.metrics_edited.get(m['name'], 0.0):.4f} | ")
            pdf.write(5, f"Human Std: {m['human']} | AI Std: {m['ai']}\n")
            
            pdf.set_font("Helvetica", "I", 10)
            pdf.multi_cell(0, 5, f"Interpretation: {m['meaning']}")
            pdf.ln(5)

    @classmethod
    def generate_excel(cls, analysis: AnalysisResult) -> bytes:
        """Generate a basic CSV/Excel compatible export."""
        import csv
        from io import StringIO
        
        output = StringIO()
        writer = csv.writer(output)
        writer.writerow(["Metric", "AI Source", "Writer Rewrite", "Delta", "Verdict"])
        
        for name, original in analysis.metrics_original.items():
            edited = analysis.metrics_edited.get(name, 0.0)
            verdict_str = "N/A"
            for res in analysis.metric_results.values():
                if res.name == name:
                    verdict_str = res.verdict
                    break
            writer.writerow([name, original, edited, edited - original, verdict_str])
            
        return output.getvalue().encode("utf-8")
